import sys

import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE

def change_to_dir():
    if (len(sys.argv) < 3):
        print("Please call the program in this format: programName pathToInputFile outputFileName")
    else:
        # print(sys.argv[1])
        # string = sys.argv[1].split("\\")
        # input_file_name = string[-1]
        # string.pop(-1)
        # print(str(string))
        # # os.chdir(sys.argv[1])
        # return input_file_name
        print(sys.argv[1])
        index_for_snipping = sys.argv[1].rfind("\\")
        input_file_name = sys.argv[1][index_for_snipping+1:]
        print(input_file_name)
        # input_file_name = string[-1]
        # string.pop(-1)
        # print(str(string))
        os.chdir(sys.argv[1][:index_for_snipping])
        print(os.getcwd())
        return input_file_name

def insert_to_doc():


    document = Document('Resume.docx')

    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(1)
    obj_font.name = 'Times New Roman'
    obj_font.color.rgb = RGBColor(255,255,255)

    paragraph = document.add_paragraph().add_run('HELLO WORLD', style='CommentsStyle')

    document.save('WellerResume.docx')

if __name__ == "__main__":
    change_to_dir()